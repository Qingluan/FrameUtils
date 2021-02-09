# -*- encoding: utf-8 -*-
import xlrd
import xlwt
from xlutils.copy import copy
import datetime
import docx
import re
import os
import tempfile
import time
import logging
from termcolor import colored

from docx.shared import Inches
from copy import deepcopy
from functools import partial
from contextlib import contextmanager


def Err(*x):
    logging.error(colored("[Error]:", "red") + " ".join([str(i) for i in x]))
def Ok(*x):
    logging.info(colored("[ok]:", "green", attrs=['bold']) + " ".join([str(i) for i in x]))

def Log(*x):
    logging.info(colored("[+]"  , "blue", attrs=['bold']) + " ".join([str(i) for i in x]))


try:
    from PIL import Image
except Exception as e:
    Err("Import PIL error!, try pip install pillow")

KEY_RE = re.compile(r'(\$\{[\w\:]+?\})')


class TaskExcel:

    def __init__(self, file_path:str, key_row=0, col_limit=30):
        self._f = file_path
        self.tp = "excel"
        self.cases = []
        self.keys = []
        self.keys_smart = []
        
        self.refresh(file_path, key_row=key_row, col_limit=col_limit)
        


    def refresh(self, f, key_row=0, col_limit=30):
        data = xlrd.open_workbook(f)
        table = data.sheets()[0]
        self.tp = "excel"
        self.cases = []
        self.keys = []
        for i in range(table.nrows):
            if i == key_row:
                # extract key:
                for no,cell_val in enumerate(table.row(i)):
                    if no < col_limit and str(cell_val.value).strip() != "":
                        key_raw = str(cell_val.value)
                        if "//" in key_raw:
                            self.keys_smart.append(key_raw.split("//")[0].strip())
                        else:
                            self.keys_smart.append(key_raw.strip())
                        self.keys.append(key_raw)

            if len(self.keys) == 0:
                continue
            d = {}
            for co in range(len(self.keys)):
                vv = table.cell(i,co)
                if vv.ctype == 3:
                    val = datetime.datetime(*xlrd.xldate_as_tuple(vv.value, 0)).ctime()
                else:
                    val = vv.value
                d[self.keys[co]] = val

            self.cases.append(d)
    
    def delete(self, rows):
        if 0 in rows :return
        old_book = xlrd.open_workbook(self._f)
        table = old_book.sheet_by_index(0)
        datas = [table.row(i) for i in range(table.nrows)]
        ded = [datas[row] for row in rows]
        [datas.remove(d) for d in ded]
        excel = xlwt.Workbook(encoding='utf-8')
        excel_table = excel.add_sheet(table.name)
        for no,da in enumerate(datas):
            for col in range(len(da)):
                excel_table.write(no, col, da[col].value)
        excel.save(self._f)

    def delete_by_data(self, datas):
        D = []
        for no,cas in enumerate(self.cases):
            for d in datas:
                min_key = set(d.keys()) & set(cas.keys())
                if set([str(d[k]).strip() for k in min_key]) == set(str(cas[k]).strip() for k in min_key):
                    D.append(no+1)
        Log("delete:",D)
        
        self.delete(D)

    def merge_by(self, datas, column=0, new_file=None):
        K = self.keys[column]
        assert len(datas) > 0
        
        assert len(self.keys) > 0
        old_book = xlrd.open_workbook(self._f)
        header = self.keys
        # excel = xlwt.Workbook(encoding = 'utf-8')
        excel = copy(wb=old_book)
        old_table = old_book.sheet_by_index(0)
        excel_table = excel.get_sheet(0)
        
        sample_data = datas[0]
        if len(sample_data) > len(self.keys):
            for k in sample_data:
                if k.strip() not in self.keys:
                    k_col = len(header)
                    excel_table.write(1, k_col, k.strip())
                    self.keys.append(k.strip())
                    
                    Ok("add new header col:", k)


        merged_rows = []
        for row in range(1,old_table.nrows):
            found = False
            for no,data in enumerate(datas):
                v = data[K]
                Log("line:",row, "Data:", old_table.row(row))
                c = column
                try:
                    if str(v).strip() == str(old_table.cell(row, c).value).strip():
                        merged_rows.append((row,no, list(data.values()) ))
                        print(K,row, no, v)
                        found = True
                        break
                except IndexError:
                    continue
                if found:
                    found = False
                    break
            
        for row,no,vals in merged_rows:
            data = datas[no]
            for col in range(len(data)):
                if col >= len(header):continue
                
                val = vals[col]
                if val != '':
                    excel_table.write(row, col, val)
        if not new_file:
            excel.save(self._f)
        else:
            if isinstance(new_file, list):
                for f in new_file:
                    excel.save(f)
            else:
                excel.save(f)
        

    def is_fulled(self, row=None, **filter):
        if row and row < len(self.cases):
            row_data = self.cases[row]
            for v in row_data.values():
                if str(v).strip() == '':
                    return False
            return True
        
        for kk,vvv in filter.items():
            if kk in self.keys:
                for case in self.cases:    
                    if case[kk] == vvv:
                        for v in case.values():
                            if str(v).strip() == "":
                                return False
                        return True
        return None



    def merge_to(self, new_file, rows):
        rows = [int(i) for i in rows]
        old_book = xlrd.open_workbook(new_file)
        datas = []
        for r in rows:
            datas.append(old_book.sheet_by_index(0).row(r))
        header = self.keys
        # excel = xlwt.Workbook(encoding = 'utf-8')
        excel = copy(wb=old_book)
        excel_table = excel.add_sheet("merged sheet %d" % (len(old_book.sheets()) + 1 ))
        for no,k in enumerate(header):
            excel_table.write(0, no, k)
        for row in range(1,len(datas)+1):
            data = datas[row-1]
            for n,v in enumerate(data):
                excel_table.write(row, n, v.value)
        excel.save(new_file)
            
    def render_table(self, redner_func, default="case.html", padding_key=None, is_admin=False):
        # if secrity:
        #     for case in self.cases:
        #         fulled = True
        #         for v in case.values():
        #             if v == "":
        #                 fulled = False
        #                 break
        if padding_key != None:
            if padding_key not in self.keys:
                self.keys.append(padding_key)
            for case in self.cases:
                if padding_key not in case:
                    case[padding_key] = ''
        return redner_func(default, cases=self.cases,enumerate=enumerate, is_admin=is_admin)
    
    def parse_headers(self, key):
        if '//' not in key:
            return {
                "name": key.strip(),
                "type": "text",
            }
        else:
            key_name, key_body = key.split("//",1)
            dd = {"name": key_name.strip()}
            if "/" in key_body:
                opts = key_body.split("/")
                dd['type'] = 'select'
                dd["options"] = opts
            else:
                dd['type'] = key_body.strip()
            
            return dd


    def render_form(self, render_func, template_name="default_form.html", action="/excel", **kargs):
        return render_func(template_name, forms={
            k: self.parse_headers(k) for k in self.keys if k.strip() != ""},
            action=action,
            hiddens=kargs,
            enumerate=enumerate,
        )
    
    @contextmanager
    def parse_from_flask_request(self, request):
        form = request.form
        try:
            d = {}
            K = self.keys_smart if self.tp == "excel" else self.keys
            for k in K:
                if k.strip() == "":continue
                if k in form:
                    d[k] = form[k]
            d['num'] = form['num']
            yield d
        except Exception as e:

            raise e
        finally:
            pass

    def restore(self,saved_name, data, padding_key=None):
        t = set(self.keys)
        if self.tp == "excel":
            t = set(self.keys_smart)
        if 'num' not in t:
            t.add("num")
        if '填写人' not in t:
            t.add('填写人')
        if set(data.keys()) != t :
            logging.error(colored("[merge excel error]:", "red"), "not same table", data.keys(), "\n",t)
            raise Exception("merge excel error")
    
        old_book = xlrd.open_workbook(saved_name)
        # add in below
        if data['num'] == '-1':
            nrows = old_book.sheets()[0].nrows
        # edit in some row
        else:
            nrows = int(data['num'])
        del data['num']
        header = old_book.sheets()[0].row(0)
        no_need_padding = False
        if padding_key != None:
            for v in header:
                if v.value == padding_key:
                    no_need_padding = True
                    break
        
        Ok("insert to => ", nrows)
        excel = copy(wb=old_book)
        excel_table = excel.get_sheet(0)
        if not no_need_padding and padding_key != None:
            excel_table.write(0, len(header), padding_key)
        for n,v in enumerate(data.values()):
            excel_table.write(nrows, n, v)
        excel.save(saved_name)
        return True


class TaskWord:
    def __init__(self, file_path):
        self._doc = docx.Document(file_path)
        self.tp = "word"
        self._keys = []
        self._reverse = {}
        self._map = {}
        self._old_doc = deepcopy(self._doc)
        self._data = {}
        self.scan()

    def keys(self):
        return self._keys
    
    def scan(self):
        paragraphs = []
        self._keys = []
        self._map = {}
        for t in self._doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraphs.append(paragraph)
                        if '${' in paragraph.text:
                            for k in KEY_RE.findall(paragraph.text):
                                self._keys.append(k)


        for paragraph in self._doc.paragraphs:
            paragraphs.append(paragraph)
            if '${' in paragraph.text:
                for k in KEY_RE.findall(paragraph.text):
                    self._keys.append(k)
        
        for k in self._keys:
            d = {}
            if ':' in k:
                label, name = k.split(":",1)
                label = label[2:]
                name = name[:-1]
            else:
                label = "textarea"
                name = k[2:-1]
            d["name"] = name
            d["type"] = label
            self._map[k] = d
            self._reverse[name] = k
        
        return paragraphs

    def restore(self, saved_name, data,smart_newline=True, padding_key=None):
        """
        :param saved_name: saved file path
        :param data: a dict must include in self._reverse.
            used by :    ... with  w.parse_from_flask_request(request) as data:
                         ...    w.docx_render(some_path, data)
        """
        if len(data) == 0:return
        data = data
        paragraphs = self.scan()

        for p in paragraphs:
            for key, val in data.items():
                key_name = '${{{}}}'.format(key) # I'm using placeholders in the form ${PlaceholderName}
                
                if key_name in p.text:
                    inline = p.runs
                    # Replace strings and retain the same style.
                    # The text to be replaced can be split over several runs so
                    # search through, identify which runs need to have text replaced
                    # then replace the text in those identified
                    started = False
                    key_index = 0
                    # found_runs is a list of (inline index, index of match, length of match)
                    found_runs = list()
                    found_all = False
                    replace_done = False

                    found_num = 0
                    # insert_image = False
                    if key_name.startswith("${img:"):
                        img_run = p.add_run()
                        if isinstance(val, list):
                            for v in val:    
                                if os.path.exists(v) and v.split(".")[-1] in ('png','jpg','jpeg'):
                                    try:
                                        img_run.add_picture(v, width=Inches(4.0))
                                    except docx.image.exceptions.UnrecognizedImageError:
                                    
                                        Err("add picture failed:", v, "try convert to png")
                                        img = Image.open(v)
                                        vpng = v.rsplit(".",1)[0] + ".png"
                                        img.save(vpng)
                                        Log("try add_picture by png...")
                                        img_run.add_picture(vpng, width=Inches(4.0))
                        # insert_image = True
                        val = ""
                    if isinstance(val, str) and smart_newline and "\n" in val:
                        
                        val = val.replace("\n", "\n\t")
                        if '\r\n' in val:
                            val = val.replace('\r\n', '\n')

                    for i in range(len(inline)):
                        Log("run:", inline[i].text,"|", key_name, key_index)
                        # case 1: found in single run so short circuit the replace
                        if key_name in inline[i].text and not started:
                            found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                            text = inline[i].text.replace(key_name, str(val))
                            inline[i].text = text
                            replace_done = True
                            found_all = True
                            break


                        # case for ${text:key} :
                        #   found HEAD
                        if "$" in inline[i].text:
                            start_index = inline[i].text.rfind("$")
                            found_left_part = inline[i].text[start_index:]
                            found_num +=1
                            if found_left_part in key_name:
                                key_index += len(found_left_part)
                                found_runs.append((i, start_index, len(found_left_part)))
                                started = True
                            if key_index == len(key_name):
                                found_all = True
                                started = False
                                key_index = 0
                                found_num = 0
                            else:
                                continue
                        
                        
                        # case for ${text:key} :
                        #    found MIDLE
                        if started and inline[i].text in key_name[key_index:]:
                            start_index = 0
                            found_left_part = inline[i].text
                            found_runs.append((i, start_index, len(found_left_part)))
                            key_index += len(found_left_part)
                            
                            found_num +=1
                            Ok("MIDDLE:", key_index)
                            if key_index == len(key_name):
                                found_all = True
                                started = False
                                key_index = 0
                                found_num = 0
                            else:
                                continue
                        
                        # case for ${text:key} :
                        #   found END
                        if started and  inline[i].text.startswith(key_name[key_index:]):
                            start_index = 0
                            found_left_part = key_name[key_index:]
                            found_runs.append((i, start_index, len(found_left_part)))
                            key_index += len(found_left_part)
                            
                            Ok("END:", key_index)
                            if key_index == len(key_name):
                                found_all = True
                                started = False
                                key_index = 0
                                found_num = 0
                            else:
                                continue

                        # case for ${text:key} :
                        #  started but not found MIDDLE , reset key_index 
                        if started and inline[i].text not in key_name[key_index:]:
                            started = False
                            key_index = 0
                            for i in range(found_num):
                                found_runs.pop()
                            found_num = 0
                            # after reset : try match first.
                            if "$" in inline[i].text:
                                start_index = inline[i].text.rfind("$")
                                found_left_part = inline[i].text[start_index:]
                                
                                if found_left_part in key_name:
                                    key_index += len(found_left_part)
                                    found_runs.append((i, start_index, len(found_left_part)))
                                    started = True
                                if key_index == len(key_name):
                                    found_all = True
                                    started = False
                                    key_index = 0
                                else:
                                    continue
                            


                    if found_all and not replace_done:
                        start_count = True
                        replace_one = 0
                        print(found_runs)
                        for i, item in enumerate(found_runs):
                            index, start, length = [t for t in item]
                            replace_one += length
                            if start_count:
                                text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                                inline[index].text = text
                                start_count = False
                            else:
                                
                                text = inline[index].text.replace(inline[index].text[start:start + length], '')
                                inline[index].text = text
                                if replace_one == len(key_name):
                                    start_count = True
                                    replace_one = 0
                                    # print("<INIT>")
                                # continue
                            # Ok(inline[index])
                            Ok("Replaced:",replace_one,"/",len(key_name),"||", p.text)
                            # if i == 0:
                            #     text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                            #     inline[index].text = text
                           
                    

        self._doc.save(saved_name)
        self._doc = deepcopy(self._old_doc)

    def render_form(self,redner_func, template_name="default_form.html", action="/word", **kargs):
        """
        :param render_func: a render functioin like flask.render_template
        """
        for k in self._keys:
            d = {}
            if ':' in k:
                label, name = k.split(":",1)
                label = label[2:]
                name = name[:-1]
            else:
                label = "textarea"
                name = k[2:-1]
            d["name"] = name
            d["type"] = label
            self._map[k] = d
            self._reverse[name] = k
        
        return redner_func(template_name, forms=self._map, action=action, hiddens=kargs,enumerate=enumerate)
    
    @contextmanager
    def parse_from_flask_request(self, request):
        """
        :param request: inlcude request.form and request.files 
        """
        if len(self._reverse) == 0:
            self.scan()
        try:
            form = request.form
            files = request.files
            with tempfile.TemporaryDirectory() as tmpdir:
                data = {}
                reverse_map = self._reverse
                for k in files.keys():
                    if k not in reverse_map:
                        Log("not found in map:",k)
                        continue
                    f_handlers = files.getlist(k)
                    data[reverse_map[k][2:-1]] = []
                    for f_handler in f_handlers:
                        val = os.path.join(tmpdir, f_handler.filename)
                        if os.path.isdir(val):
                            continue
                        f_handler.save(val)
                        data[reverse_map[k][2:-1]].append(val)
                        Ok("cache->", val)

                for k in form.keys():
                    if k not in reverse_map:
                        Log("not found in map:",k)
                        continue
                    
                    val = form[k]
                    data[reverse_map[k][2:-1]] = val
                yield data
        except Exception as e:
            raise e
        finally:
            pass



def Office(f):
    if f.rsplit(".")[-1] in ("xlsx", "xls"):
        return TaskExcel(f)
    elif f.rsplit(".")[-1] in ("doc", "docx"):
        return TaskWord(f)
    else:
        raise Exception("not supported file to use Office!")